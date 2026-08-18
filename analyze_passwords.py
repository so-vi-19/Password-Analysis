#!/usr/bin/env python3
"""
analyze_passwords.py

Analyzes results from a password-cracking audit (e.g. hashcat/john --show output)
to report:
  - How many accounts were cracked out of the total attempted
  - How many cracked passwords contain a campus-name variation
    (handles case, leetspeak substitutions, and partial/mascot matches)
  - How many *sensitive* accounts (admins, faculty, finance, etc.) were cracked
  - A general breakdown of other weak patterns (season+year, keyboard walks,
    common base words) for context
  - General descriptive statistics about the cracked password set

Intended for legitimate security auditing / password-policy review of an
organization's own credentials (e.g. an internal red-team or IT security
report after a sanctioned cracking exercise).

------------------------------------------------------------------------
CASE-INSENSITIVITY (important)
------------------------------------------------------------------------
All identity/keyword matching in this script is case-insensitive:
  - Usernames are matched against --sensitive-accounts regardless of case
    (e.g. "Admin", "ADMIN", and "admin" are treated as the same account).
  - Campus-name / mascot keyword matching against passwords is also
    case-insensitive (handled by normalize()).
Username *display* in the report preserves the original casing from the
--cracked file; only the comparison itself is case-folded.

------------------------------------------------------------------------
INPUT FORMATS SUPPORTED  (--cracked FILE)
------------------------------------------------------------------------
By default (no --has-username), one entry per line, password-only or
hash:password — the script just takes the LAST colon-delimited field as
the password and doesn't track usernames:
    password
    hash:password

With --has-username (use this for hashcat's `--username` output, or any
john/hashcat --show export that includes the account name), the FIRST
field is treated as the username and the LAST field as the password.
Anything in between is treated as the hash:
    username:password
    username:hash:password
    username:hash:hash:password   (some hash formats embed extra colons)

------------------------------------------------------------------------
--total FILE / --total-count N (optional)
------------------------------------------------------------------------
--total FILE is a file with one username/hash per line representing every
account attempted (cracked + uncracked), used to compute crack rate.
--total-count N is a plain integer alternative if you don't have/want to
share the full list.

------------------------------------------------------------------------
--sensitive-accounts FILE (optional, requires --has-username)
------------------------------------------------------------------------
A file listing sensitive usernames (admins, service accounts, finance,
executives, etc.), one per line, matched case-insensitively against the
usernames in --cracked. The report includes how many of these were
cracked, out of how many total sensitive accounts, plus the list of
which ones (with campus-variant flag).

------------------------------------------------------------------------
CAMPUS KEYWORDS
------------------------------------------------------------------------
--campus-name "State University"
    Full campus name. The script automatically derives useful variants:
    acronym, each significant word, and a no-space concatenation.

--keywords-file FILE (optional)
    A plain text file, one keyword/variant per line, for things a name-parser
    can't guess: mascot, city, abbreviation, athletics nickname, old name, etc.

------------------------------------------------------------------------
USAGE EXAMPLE
------------------------------------------------------------------------
python3 analyze_passwords.py \
    --cracked cracked.txt --has-username \
    --total-count 5000 \
    --campus-name "Riverside State University" \
    --keywords-file campus_keywords.txt \
    --sensitive-accounts sensitive_users.txt \
    --out report.json --csv report.csv --docx report.docx

Requires: python-docx (pip install python-docx --break-system-packages)
"""

import argparse
import csv
import json
import re
import sys
from collections import Counter
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Case-insensitive comparison helper (used for ALL username matching)
# ---------------------------------------------------------------------------

def fold_username(username: str) -> str:
    """Canonical form of a username for case-insensitive comparison:
    strip surrounding whitespace, lowercase. Always use this on BOTH sides
    of any username comparison so casing differences between the cracked
    file and the sensitive-accounts file never cause a missed match."""
    if username is None:
        return ""
    return username.strip().lower()


def username_variants(username: str) -> set:
    """Returns every reasonable case-folded form of a username so that
    formatting differences between the --cracked file and the
    --sensitive-accounts file don't cause a missed match. Covers:
      - plain case-folded username
      - DOMAIN\\username  -> username   (Windows/AD-style domain prefix)
      - username@domain   -> username   (UPN / email-style suffix)
    A match is counted if ANY variant of the cracked username matches ANY
    variant of a sensitive-list entry.
    """
    folded = fold_username(username)
    if not folded:
        return set()
    variants = {folded}
    if "\\" in folded:
        variants.add(folded.split("\\")[-1])
    if "/" in folded:
        variants.add(folded.split("/")[-1])
    if "@" in folded:
        variants.add(folded.split("@")[0])
    return variants


# ---------------------------------------------------------------------------
# Leetspeak / normalization (for password -> campus-keyword matching)
# ---------------------------------------------------------------------------

LEET_MAP = {
    "@": "a", "4": "a",
    "3": "e",
    "1": "i", "!": "i", "|": "i",
    "0": "o",
    "$": "s", "5": "s",
    "7": "t", "+": "t",
    "9": "g",
}

def normalize(password: str) -> str:
    """Lowercase and collapse common leetspeak substitutions so 'C4mpu5!'
    normalizes to 'campusi' style comparable text. Case-insensitive by
    construction (password.lower() is always applied first)."""
    pw = password.lower()
    for leet, letter in LEET_MAP.items():
        pw = pw.replace(leet, letter)
    return re.sub(r"[^a-z0-9]", "", pw)


def derive_campus_variants(campus_name: str) -> set:
    """From a full campus name, derive likely password-relevant substrings.
    All variants are lowercase since normalize() always lowercases the
    password side before comparison."""
    stopwords = {"of", "the", "at", "and", "university", "college", "state"}
    words = re.findall(r"[a-zA-Z]+", campus_name.lower())
    variants = set()

    full_concat = "".join(words)
    if full_concat:
        variants.add(full_concat)

    acronym = "".join(w[0] for w in words if w)
    if len(acronym) >= 2:
        variants.add(acronym)

    for w in words:
        if w not in stopwords and len(w) >= 3:
            variants.add(w)

    return variants


def load_keyword_set(path: str) -> set:
    variants = set()
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            line = line.strip().lower()
            if line:
                variants.add(re.sub(r"[^a-z0-9]", "", line))
    return variants


def load_sensitive_usernames(path: str) -> set:
    """Returns a set of ALL case-folded variant forms (see username_variants)
    of every username in the file, ready for case-insensitive, domain-prefix-
    insensitive, and UPN-suffix-insensitive membership checks."""
    names = set()
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            line = line.strip()
            if line:
                names |= username_variants(line)
    return names


# ---------------------------------------------------------------------------
# Other weak-pattern heuristics (for extra context in the report)
# ---------------------------------------------------------------------------

KEYBOARD_WALKS = ["qwerty", "asdf", "zxcv", "12345", "09876", "qazwsx"]
SEASON_WORDS = ["spring", "summer", "fall", "autumn", "winter"]

def detect_other_patterns(pw_norm: str) -> list:
    hits = []
    if any(walk in pw_norm for walk in KEYBOARD_WALKS):
        hits.append("keyboard_walk")
    if any(season in pw_norm for season in SEASON_WORDS) and re.search(r"(19|20)\d{2}", pw_norm):
        hits.append("season_plus_year")
    elif any(season in pw_norm for season in SEASON_WORDS):
        hits.append("season_word")
    if re.fullmatch(r"[a-z]+\d{2,4}", pw_norm):
        hits.append("word_plus_trailing_digits")
    if re.search(r"(19|20)\d{2}", pw_norm):
        hits.append("contains_year")
    if re.fullmatch(r"(.)\1{3,}", pw_norm):
        hits.append("repeated_char")
    return hits


# ---------------------------------------------------------------------------
# Parsing cracked-password lines
# ---------------------------------------------------------------------------

def parse_line(line: str, has_username: bool):
    """Returns (username_or_None, password) for one line of cracked output.

    has_username=False (default): last colon-delimited field is the password,
        no username tracked (e.g. plain password list, or hash:password).
    has_username=True: first field is username, last field is password,
        anything in between is the hash.
    """
    line = line.rstrip("\n").rstrip("\r")
    if not line:
        return None, ""

    if not has_username:
        if ":" in line:
            return None, line.split(":")[-1]
        return None, line

    parts = line.split(":")
    if len(parts) == 1:
        # No colon at all even though --has-username was set; treat whole
        # line as password with no recoverable username.
        return None, parts[0]
    username = parts[0].strip()
    password = parts[-1]
    return username, password


def load_cracked(path: str, has_username: bool) -> list:
    """Returns a list of dicts: {"username": str_or_None, "password": str}"""
    entries = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            username, pw = parse_line(line, has_username)
            if pw != "":
                entries.append({"username": username, "password": pw})
    return entries


def count_lines(path: str) -> int:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return sum(1 for line in f if line.strip())


# ---------------------------------------------------------------------------
# General password statistics (independent of campus/sensitive matching)
# ---------------------------------------------------------------------------

def compute_general_stats(cracked_entries, top_n=10):
    passwords = [e["password"] for e in cracked_entries]
    n = len(passwords)
    if n == 0:
        return {"total_cracked": 0}

    lengths = [len(pw) for pw in passwords]
    avg_len = round(sum(lengths) / n, 2)

    buckets = Counter()
    for l in lengths:
        if l < 8:
            buckets["<8"] += 1
        elif l <= 9:
            buckets["8-9"] += 1
        elif l <= 11:
            buckets["10-11"] += 1
        elif l <= 15:
            buckets["12-15"] += 1
        else:
            buckets["16+"] += 1
    length_distribution = {k: buckets.get(k, 0) for k in ["<8", "8-9", "10-11", "12-15", "16+"]}

    has_upper = sum(1 for pw in passwords if re.search(r"[A-Z]", pw))
    has_lower = sum(1 for pw in passwords if re.search(r"[a-z]", pw))
    has_digit = sum(1 for pw in passwords if re.search(r"\d", pw))
    has_symbol = sum(1 for pw in passwords if re.search(r"[^A-Za-z0-9]", pw))
    digits_only = sum(1 for pw in passwords if re.fullmatch(r"\d+", pw))
    lower_only = sum(1 for pw in passwords if re.fullmatch(r"[a-z]+", pw))
    alpha_only = sum(1 for pw in passwords if re.fullmatch(r"[A-Za-z]+", pw))

    def pct(x):
        return round(100 * x / n, 2)

    composition = {
        "has_uppercase_pct": pct(has_upper),
        "has_lowercase_pct": pct(has_lower),
        "has_digit_pct": pct(has_digit),
        "has_symbol_pct": pct(has_symbol),
        "digits_only_pct": pct(digits_only),
        "lowercase_only_pct": pct(lower_only),
        "letters_only_pct": pct(alpha_only),
    }

    freq = Counter(passwords)
    unique_count = len(freq)
    reused = {pw: c for pw, c in freq.items() if c > 1}
    accounts_sharing_passwords = sum(reused.values())

    top_common = [{"password": pw, "count": c} for pw, c in freq.most_common(top_n)]
    top_reused = [{"password": pw, "count": c} for pw, c in
                  sorted(reused.items(), key=lambda x: -x[1])[:top_n]]

    return {
        "total_cracked": n,
        "unique_passwords": unique_count,
        "duplicate_unique_passwords": len(reused),
        "accounts_sharing_a_password": accounts_sharing_passwords,
        "avg_length": avg_len,
        "min_length": min(lengths),
        "max_length": max(lengths),
        "length_distribution": length_distribution,
        "composition": composition,
        "top_common_passwords": top_common,
        "top_reused_passwords": top_reused,
    }


def print_general_stats(stats):
    if stats.get("total_cracked", 0) == 0:
        return
    print("\n" + "-" * 60)
    print("GENERAL PASSWORD STATISTICS")
    print("-" * 60)
    print(f"Unique passwords          : {stats['unique_passwords']} of {stats['total_cracked']}")
    print(f"Accounts sharing a password: {stats['accounts_sharing_a_password']} "
          f"(across {stats['duplicate_unique_passwords']} reused password value(s))")
    print(f"Length (avg / min / max)  : {stats['avg_length']} / {stats['min_length']} / {stats['max_length']}")
    print("Length distribution       :", stats["length_distribution"])
    print("Composition:")
    for k, v in stats["composition"].items():
        print(f"  {k:<22} {v}%")
    if stats["top_common_passwords"]:
        print("Top common passwords:")
        for item in stats["top_common_passwords"]:
            print(f"  {item['password']:<20} x{item['count']}")


# ---------------------------------------------------------------------------
# Main analysis
# ---------------------------------------------------------------------------

def analyze(cracked_entries, total_count, campus_variants, sensitive_usernames, total_sensitive_count):
    """
    sensitive_usernames: a set of CASE-FOLDED usernames (see fold_username),
    or None if --sensitive-accounts wasn't provided.

    Sensitive-account matching always case-folds the cracked entry's
    username before checking membership, so "Admin", "ADMIN", "admin",
    and " admin " (stray whitespace) are all recognized as the same
    sensitive account regardless of how they were capitalized in either
    the --cracked file or the --sensitive-accounts file.
    """
    n_cracked = len(cracked_entries)
    campus_matches = []
    other_pattern_counter = Counter()
    matched_variant_counter = Counter()
    sensitive_cracked = []
    seen_sensitive = set()  # avoid double-counting if the same account appears twice

    for entry in cracked_entries:
        pw = entry["password"]
        username = entry["username"]
        pw_norm = normalize(pw)

        matched_variants = [v for v in campus_variants if v and v in pw_norm]
        is_campus_match = bool(matched_variants)
        if is_campus_match:
            campus_matches.append({"username": username, "password": pw, "matched_on": matched_variants})
            for v in matched_variants:
                matched_variant_counter[v] += 1

        for pattern in detect_other_patterns(pw_norm):
            other_pattern_counter[pattern] += 1

        if sensitive_usernames is not None and username is not None:
            candidate_variants = username_variants(username)
            match_key = next(iter(candidate_variants & sensitive_usernames), None) if candidate_variants else None
            if match_key and match_key not in seen_sensitive:
                seen_sensitive.add(match_key)
                sensitive_cracked.append({
                    "username": username,  # original casing preserved for display
                    "password": pw,
                    "campus_variation": is_campus_match,
                })

    n_campus = len(campus_matches)

    report = {
        "total_accounts_attempted": total_count,
        "total_cracked": n_cracked,
        "crack_rate_pct": round(100 * n_cracked / total_count, 2) if total_count else None,
        "campus_variation_count": n_campus,
        "campus_variation_pct_of_cracked": round(100 * n_campus / n_cracked, 2) if n_cracked else None,
        "campus_variation_pct_of_total": round(100 * n_campus / total_count, 2) if total_count else None,
        "matched_variant_frequency": dict(matched_variant_counter.most_common()),
        "other_weak_patterns": dict(other_pattern_counter.most_common()),
        "campus_matches_sample": campus_matches[:50],
    }

    if sensitive_usernames is not None:
        n_sensitive_total = total_sensitive_count if total_sensitive_count is not None else len(sensitive_usernames)
        n_sensitive_cracked = len(sensitive_cracked)
        report["sensitive_accounts_total"] = n_sensitive_total
        report["sensitive_accounts_cracked"] = n_sensitive_cracked
        report["sensitive_accounts_cracked_pct"] = (
            round(100 * n_sensitive_cracked / n_sensitive_total, 2) if n_sensitive_total else None
        )
        report["sensitive_accounts_cracked_list"] = sensitive_cracked

    report["general_stats"] = compute_general_stats(cracked_entries)

    return report, campus_matches


def print_summary(report):
    print("=" * 60)
    print("PASSWORD CRACK ANALYSIS SUMMARY")
    print("=" * 60)
    if report["total_accounts_attempted"]:
        print(f"Total accounts attempted : {report['total_accounts_attempted']}")
        print(f"Total cracked             : {report['total_cracked']} "
              f"({report['crack_rate_pct']}%)")
    else:
        print(f"Total cracked             : {report['total_cracked']} "
              f"(total attempted not provided)")

    print(f"Cracked using campus-name variation : {report['campus_variation_count']} "
          f"({report['campus_variation_pct_of_cracked']}% of cracked"
          + (f", {report['campus_variation_pct_of_total']}% of all accounts)"
             if report["campus_variation_pct_of_total"] is not None else ")"))

    if "sensitive_accounts_total" in report:
        print(f"\nSensitive accounts cracked : {report['sensitive_accounts_cracked']} "
              f"of {report['sensitive_accounts_total']} "
              f"({report['sensitive_accounts_cracked_pct']}%)")
        if report["sensitive_accounts_cracked_list"]:
            print("  Cracked sensitive accounts:")
            for entry in report["sensitive_accounts_cracked_list"]:
                flag = " [campus-variant]" if entry["campus_variation"] else ""
                print(f"    {entry['username']:<20} {entry['password']}{flag}")

    if report["matched_variant_frequency"]:
        print("\nMost common campus-variant matches:")
        for variant, count in report["matched_variant_frequency"].items():
            print(f"  {variant:<20} {count}")

    if report["other_weak_patterns"]:
        print("\nOther weak patterns observed (context, not campus-specific):")
        for pattern, count in report["other_weak_patterns"].items():
            print(f"  {pattern:<28} {count}")
    print("=" * 60)

    if "general_stats" in report:
        print_general_stats(report["general_stats"])


def write_csv(campus_matches, path):
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["username", "password", "matched_on"])
        for m in campus_matches:
            writer.writerow([m.get("username") or "", m["password"], ";".join(m["matched_on"])])


# ---------------------------------------------------------------------------
# DOCX report generation (python-docx — no external Node/JS dependency)
# ---------------------------------------------------------------------------

HEADER_FILL = "1F3864"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
MUTED_TEXT = RGBColor(0x66, 0x66, 0x66)


def _shade_cell(cell, hex_color):
    """Apply background shading to a table cell (python-docx has no
    built-in shading API, so we drop into the underlying XML)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _add_muted(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED_TEXT
    run.font.size = Pt(9)
    return p


def _add_table(doc, header_row, rows):
    """Adds a simple header+data table styled with a dark header band."""
    table = doc.add_table(rows=1, cols=len(header_row))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header_row):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(str(text))
        run.bold = True
        run.font.color.rgb = HEADER_TEXT
        _shade_cell(hdr_cells[i], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()  # spacing after table
    return table


def write_docx(report, docx_path, campus_name):
    """Builds a formatted Word report directly from the report dict using
    python-docx."""
    doc = Document()
    campus_label = campus_name or "Campus"

    doc.add_heading(f"{campus_label} — Password Crack Analysis Report", level=0)
    _add_muted(doc, f"Generated {date.today().isoformat()}")

    # --- Overview ---------------------------------------------------------
    _add_heading(doc, "Overview", level=1)
    overview_rows = []
    if report.get("total_accounts_attempted"):
        overview_rows.append(["Total accounts attempted", str(report["total_accounts_attempted"])])
    overview_rows.append(["Total cracked", str(report["total_cracked"])])
    if report.get("crack_rate_pct") is not None:
        overview_rows.append(["Crack rate", f"{report['crack_rate_pct']}%"])
    overview_rows.append(["Cracked using a campus-name variation", str(report["campus_variation_count"])])
    if report.get("campus_variation_pct_of_cracked") is not None:
        overview_rows.append(["  ...as % of cracked passwords", f"{report['campus_variation_pct_of_cracked']}%"])
    if report.get("campus_variation_pct_of_total") is not None:
        overview_rows.append(["  ...as % of all attempted accounts", f"{report['campus_variation_pct_of_total']}%"])
    if "sensitive_accounts_total" in report:
        overview_rows.append([
            "Sensitive accounts cracked",
            f"{report['sensitive_accounts_cracked']} of {report['sensitive_accounts_total']} "
            f"({report['sensitive_accounts_cracked_pct']}%)",
        ])
    _add_table(doc, ["Metric", "Value"], overview_rows)

    # --- General password statistics ---------------------------------------
    gs = report.get("general_stats", {})
    if gs.get("total_cracked"):
        _add_heading(doc, "General Password Statistics", level=1)
        gen_rows = [
            ["Unique passwords", f"{gs['unique_passwords']} of {gs['total_cracked']}"],
            ["Accounts sharing a password with another account",
             f"{gs['accounts_sharing_a_password']} (across {gs['duplicate_unique_passwords']} reused value(s))"],
            ["Average length", str(gs["avg_length"])],
            ["Shortest / longest password", f"{gs['min_length']} / {gs['max_length']} characters"],
        ]
        _add_table(doc, ["Metric", "Value"], gen_rows)

        _add_heading(doc, "Length Distribution", level=2)
        len_rows = [[bucket, str(count)] for bucket, count in gs.get("length_distribution", {}).items()]
        _add_table(doc, ["Length (chars)", "Count"], len_rows)

        _add_heading(doc, "Character Composition", level=2)
        comp_labels = {
            "has_uppercase_pct": "Contains an uppercase letter",
            "has_lowercase_pct": "Contains a lowercase letter",
            "has_digit_pct": "Contains a digit",
            "has_symbol_pct": "Contains a symbol",
            "digits_only_pct": "Digits only",
            "lowercase_only_pct": "Lowercase letters only",
            "letters_only_pct": "Letters only (no digits/symbols)",
        }
        comp_rows = [[comp_labels.get(k, k), f"{v}%"] for k, v in gs.get("composition", {}).items()]
        _add_table(doc, ["Property", "% of cracked passwords"], comp_rows)

        if gs.get("top_common_passwords"):
            _add_heading(doc, "Most Common Cracked Passwords", level=2)
            top_rows = [[item["password"], str(item["count"])] for item in gs["top_common_passwords"]]
            _add_table(doc, ["Password", "Times used"], top_rows)

    # --- Campus name variation ----------------------------------------------
    _add_heading(doc, "Campus-Name Variation in Passwords", level=1)
    pct_note = (
        f" — {report['campus_variation_pct_of_cracked']}% of all cracked passwords."
        if report.get("campus_variation_pct_of_cracked") is not None else "."
    )
    doc.add_paragraph(
        f"{report['campus_variation_count']} cracked password(s) contained a variation of the "
        f"campus name, mascot, or related keyword{pct_note}"
    )

    variant_freq = report.get("matched_variant_frequency", {})
    if variant_freq:
        _add_heading(doc, "Matched Keyword Frequency", level=2)
        variant_rows = [[variant, str(count)] for variant, count in variant_freq.items()]
        _add_table(doc, ["Matched keyword", "Occurrences"], variant_rows)

    campus_sample = report.get("campus_matches_sample", [])
    if campus_sample:
        _add_heading(doc, "Cracked Passwords Matching Campus Keywords", level=2)
        rows = [[m.get("username") or "(n/a)", m["password"], ", ".join(m.get("matched_on", []))]
                for m in campus_sample]
        _add_table(doc, ["Username", "Password", "Matched on"], rows)
        if len(campus_sample) < report["campus_variation_count"]:
            _add_muted(
                doc,
                f"Showing {len(campus_sample)} of {report['campus_variation_count']} matches. "
                "See the accompanying CSV for the full list.",
            )

    # --- Sensitive accounts ---------------------------------------------------
    if "sensitive_accounts_total" in report:
        _add_heading(doc, "Sensitive Account Exposure", level=1)
        doc.add_paragraph(
            f"{report['sensitive_accounts_cracked']} of {report['sensitive_accounts_total']} "
            f"sensitive accounts were cracked ({report['sensitive_accounts_cracked_pct']}%)."
        )
        sens_list = report.get("sensitive_accounts_cracked_list", [])
        if sens_list:
            rows = [[s["username"], s["password"], "Yes" if s["campus_variation"] else "No"] for s in sens_list]
            _add_table(doc, ["Username", "Password", "Campus-name variant?"], rows)
        else:
            doc.add_paragraph("No sensitive accounts were among the cracked passwords.")

    # --- Other weak patterns -----------------------------------------------
    other_patterns = report.get("other_weak_patterns", {})
    if other_patterns:
        _add_heading(doc, "Other Weak Patterns Observed", level=1)
        doc.add_paragraph("Additional patterns noted for context (not campus-specific):")
        rows = [[pattern.replace("_", " "), str(count)] for pattern, count in other_patterns.items()]
        _add_table(doc, ["Pattern", "Count"], rows)

    doc.save(docx_path)
    return True


def print_debug_diagnostics(cracked_entries, sensitive_accounts_path, sensitive_usernames):
    """Prints exactly what the script parsed, so username-format mismatches
    (wrong delimiter, missing --has-username, domain prefixes it didn't
    expect, etc.) are visible instead of silently producing 0 matches."""
    print("\n" + "#" * 60)
    print("DEBUG DIAGNOSTICS")
    print("#" * 60)

    total = len(cracked_entries)
    with_username = sum(1 for e in cracked_entries if e["username"])
    print(f"Parsed {total} cracked entries; {with_username} have a non-empty username "
          f"({total - with_username} do NOT — check --has-username / your file's delimiter).")

    print("\nFirst 5 parsed entries (username -> password):")
    for e in cracked_entries[:5]:
        print(f"  {e['username']!r:25} -> {e['password']!r}")

    if sensitive_accounts_path:
        print(f"\nSensitive-accounts file: {sensitive_accounts_path}")
        print(f"Loaded {len(sensitive_usernames)} variant form(s) after case/domain/UPN normalization.")
        print("Sample normalized sensitive variants:", sorted(list(sensitive_usernames))[:15])

        cracked_usernames = sorted({e["username"] for e in cracked_entries if e["username"]})
        print(f"\n{len(cracked_usernames)} unique username(s) found in --cracked. Sample:")
        print(" ", cracked_usernames[:15])

        # Which cracked usernames matched a sensitive variant, and which didn't
        unmatched = []
        for u in cracked_usernames:
            if not (username_variants(u) & sensitive_usernames):
                unmatched.append(u)
        matched = [u for u in cracked_usernames if u not in unmatched]
        print(f"\nOf the cracked usernames, {len(matched)} matched a sensitive-account variant.")
        if matched:
            print("  Matched:", matched[:20])
        print(f"({len(unmatched)} cracked usernames did NOT match any sensitive entry — expected for most accounts.)")

    print("#" * 60 + "\n")


def main():
    parser = argparse.ArgumentParser(
        description="Analyze cracked-password results for campus-name reuse, sensitive-account exposure, and other weak patterns."
    )
    parser.add_argument("--cracked", required=True, help="File of cracked passwords/hashcat output")
    parser.add_argument("--has-username", action="store_true",
                         help="Set this if --cracked lines start with username (e.g. hashcat --username output: user:hash:password)")
    parser.add_argument("--total", help="File listing every account attempted (for crack-rate calc)")
    parser.add_argument("--total-count", type=int, help="Total accounts attempted, as a plain integer (alternative to --total)")
    parser.add_argument("--campus-name", help='Full campus name, e.g. "Riverside State University"')
    parser.add_argument("--keywords-file", help="Extra keyword variants file (mascot, city, abbreviation, etc.), one per line")
    parser.add_argument("--sensitive-accounts", help="File of sensitive usernames (one per line, matched case-insensitively). Requires --has-username.")
    parser.add_argument("--sensitive-total-count", type=int,
                         help="Total number of sensitive accounts (defaults to the count of usernames in --sensitive-accounts)")
    parser.add_argument("--out", help="Write full JSON report to this path")
    parser.add_argument("--csv", help="Write matched campus-variant passwords to this CSV path")
    parser.add_argument("--docx", help="Write a formatted Word (.docx) report to this path")
    parser.add_argument("--debug", action="store_true",
                         help="Print diagnostic info about parsed usernames and sensitive-account matching to help troubleshoot mismatches")
    args = parser.parse_args()

    if not args.campus_name and not args.keywords_file:
        parser.error("Provide at least one of --campus-name or --keywords-file so there's something to match against.")

    if args.sensitive_accounts and not args.has_username:
        parser.error("--sensitive-accounts requires --has-username so usernames can be extracted from --cracked.")

    cracked_entries = load_cracked(args.cracked, args.has_username)

    total_count = None
    if args.total:
        total_count = count_lines(args.total)
    elif args.total_count:
        total_count = args.total_count

    campus_variants = set()
    if args.campus_name:
        campus_variants |= derive_campus_variants(args.campus_name)
    if args.keywords_file:
        campus_variants |= load_keyword_set(args.keywords_file)

    if not campus_variants:
        print("No usable campus keywords derived — check --campus-name / --keywords-file input.", file=sys.stderr)
        sys.exit(1)

    sensitive_usernames = None
    if args.sensitive_accounts:
        sensitive_usernames = load_sensitive_usernames(args.sensitive_accounts)

    if args.debug:
        print_debug_diagnostics(cracked_entries, args.sensitive_accounts, sensitive_usernames)

    report, campus_matches = analyze(
        cracked_entries, total_count, campus_variants,
        sensitive_usernames, args.sensitive_total_count,
    )
    print_summary(report)
    print(f"\n(Matching against {len(campus_variants)} campus variant(s): {sorted(campus_variants)})")

    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=2)
        print(f"\nFull JSON report written to {args.out}")

    if args.csv:
        write_csv(campus_matches, args.csv)
        print(f"Matched passwords written to {args.csv}")

    if args.docx:
        write_docx(report, args.docx, args.campus_name)
        print(f"Word report written to {args.docx}")


if __name__ == "__main__":
    main()
